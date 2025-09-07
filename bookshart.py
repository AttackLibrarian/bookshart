from docx import Document
from ebooklib import epub
import os

def bookshart(docx_path, epub_path, title="Untitled", author="Anonymous"):
    doc = Document(docx_path)
    book = epub.EpubBook()
    book.set_identifier("id123456")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)

    chapters = []
    current_chapter = None
    chapter_content = ""
    image_count = 0

    def flush_chapter(title, content, idx):
        c = epub.EpubHtml(title=title, file_name=f"chap_{idx}.xhtml", lang="en")
        c.content = f"<h1>{title}</h1>{content}"
        book.add_item(c)
        return c

    def parse_paragraph(para):
        html = ""
        for run in para.runs:
            text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if run.bold: text = f"<b>{text}</b>"
            if run.italic: text = f"<i>{text}</i>"
            if run.underline: text = f"<u>{text}</u>"
            html += text
        return f"<p>{html}</p>" if html.strip() else ""

    def parse_table(table):
        rows = []
        for row in table.rows:
            cells = [f"<td>{' '.join(parse_paragraph(p) for p in cell.paragraphs)}</td>" for cell in row.cells]
            rows.append(f"<tr>{''.join(cells)}</tr>")
        return f"<table>{''.join(rows)}</table>"

    for block in doc.element.body.iterchildren():
        if block.tag.endswith("p"):
            para = doc.paragraphs[len(chapters) + len(chapter_content.split("</p>"))]
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    if current_chapter:
                        chapters.append(flush_chapter(current_chapter, chapter_content, len(chapters)))
                        chapter_content = ""
                    current_chapter = para.text
                else:
                    chapter_content += parse_paragraph(para)

            for run in para.runs:
                if "graphic" in run._element.xml:
                    for rel in run.part.rels.values():
                        if "image" in rel.reltype:
                            image_count += 1
                            image_name = f"image_{image_count}.png"
                            image_data = rel.target_part.blob
                            img_item = epub.EpubItem(uid=image_name, file_name=f"images/{image_name}", media_type="image/png", content=image_data)
                            book.add_item(img_item)
                            chapter_content += f'<div><img src="images/{image_name}" alt="Image" /></div>'

        elif block.tag.endswith("tbl"):
            table = doc.tables[0]
            chapter_content += parse_table(table)

    if current_chapter:
        chapters.append(flush_chapter(current_chapter, chapter_content, len(chapters)))

    book.toc = tuple(chapters)
    book.spine = ["nav"] + chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    style = """
    body { font-family: Arial, sans-serif; line-height: 1.4; }
    h1 { margin-top: 1em; }
    p { margin: 0.5em 0; }
    table { border-collapse: collapse; width: 100%; margin: 1em 0; }
    td { border: 1px solid #444; padding: 4px; }
    img { max-width: 100%; display: block; margin: 1em auto; }
    """
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
    book.add_item(nav_css)

    epub.write_epub(epub_path, book, {})

# Example usage:
# bookshart("example.docx", "output.epub", title="BookShart Deluxe", author="atklib")
